"""Generate the ScripTree competitive analysis Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
import datetime


def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2E74B5")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D9E2F3")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def build_document():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ================================================================
    # TITLE PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ScripTree")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Competitive Analysis & Market Positioning")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph("")

    dateline = doc.add_paragraph()
    dateline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dateline.add_run(datetime.date.today().strftime("%B %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ================================================================
    # TABLE OF CONTENTS (manual)
    # ================================================================
    doc.add_heading("Contents", level=1)
    toc_items = [
        "1.  Executive Summary",
        "2.  What ScripTree Does",
        "3.  Competitive Landscape",
        "4.  Feature Comparison Matrix",
        "5.  Detailed Competitor Profiles",
        "6.  ScripTree's Unique Advantages",
        "7.  LLM-Powered Workflow",
        "8.  Market Positioning & Target Audiences",
        "9.  Mass Appeal Assessment",
        "10. Conclusion",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "ScripTree is a universal GUI generator that wraps any command-line "
        "tool\u2014regardless of language, framework, or platform\u2014in a "
        "native Windows form with file pickers, dropdowns, checkboxes, and "
        "structured output capture. Unlike every competitor surveyed, "
        "ScripTree requires zero modification to the target tool's source "
        "code and works equally well with executables that have no --help "
        "output at all."
    )
    doc.add_paragraph(
        "This document compares ScripTree against five direct competitors "
        "(Gooey, Python Fire, Clicky, Sunbeam, and Gradio) and three "
        "adjacent tools (Streamlit, Textual, and Charm/Bubbletea). It "
        "identifies ScripTree's unique position in the market and assesses "
        "its potential for broad adoption."
    )

    # ================================================================
    # 2. WHAT SCRIPTREE DOES
    # ================================================================
    doc.add_heading("2. What ScripTree Does", level=1)
    doc.add_paragraph(
        "ScripTree generates a GUI form for any CLI tool. The user points "
        "ScripTree at an executable; ScripTree either parses its --help "
        "output automatically or provides a blank-canvas editor for "
        "manual definition. Either way, the result is a .scriptree file "
        "\u2014 a portable JSON definition that describes the tool's "
        "parameters, argument template, and UI layout."
    )

    doc.add_heading("Core capabilities", level=2)
    bullets = [
        ("Auto-parse help text", "Detects argparse, Click, docopt, and "
         "generic flag patterns from --help output. Automatically assigns "
         "appropriate widget types (file pickers for paths, dropdowns for "
         "enums, checkboxes for booleans)."),
        ("Blank-canvas editor", "For tools with no --help (custom .exe, "
         "legacy tools), build the form from scratch in under 90 seconds. "
         "Full property panel with live argv preview."),
        ("Sections & tabs", "Group parameters into collapsible sections or "
         "tabbed pages for complex tools (e.g., robocopy with 60+ flags)."),
        ("Configurations", "Save and switch between named parameter presets "
         "per tool, with environment variable and PATH overrides per "
         "configuration."),
        ("Tree launcher", ".scriptreetree files group multiple tools into "
         "a navigable tree \u2014 an entire team's toolkit in one window."),
        ("Native Windows UI", "PySide6/Qt6 with native Windows dialogs for "
         "file/folder pickers. Looks and feels like a standard Windows "
         "application."),
        ("Output capture", "Real-time stdout/stderr streaming with ANSI "
         "colour support, resizable command-line panel with word wrap, "
         "and copy-to-clipboard."),
        ("LLM documentation pack", "A dedicated set of reference files "
         "enables AI assistants to generate valid .scriptree files from "
         "natural-language descriptions or --help output, without needing "
         "access to the ScripTree codebase."),
    ]
    for title_text, desc in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + ": ")
        run.bold = True
        p.add_run(desc)

    # ================================================================
    # 3. COMPETITIVE LANDSCAPE
    # ================================================================
    doc.add_heading("3. Competitive Landscape", level=1)
    doc.add_paragraph(
        "The \"CLI-to-GUI\" space is fragmented. Most tools target a single "
        "language or require source-code integration. The table below "
        "categorises every tool surveyed."
    )

    cat_headers = ["Category", "Tools", "Approach"]
    cat_rows = [
        ["Python decorator/wrapper",
         "Gooey, Clicky, Gradio",
         "Requires modifying the Python source to add decorators or "
         "wrapper calls. GUI is generated at runtime."],
        ["Python introspection",
         "Python Fire",
         "Generates a CLI (not GUI) from Python functions/classes. "
         "No visual interface."],
        ["TUI framework",
         "Textual, Charm/Bubbletea",
         "Terminal UI toolkits. Developer builds the interface manually "
         "in code. Not a generator."],
        ["Script runner / launcher",
         "Sunbeam",
         "Runs scripts that emit JSON to describe a UI. Requires each "
         "tool to be written as a Sunbeam-compatible script."],
        ["Universal GUI generator",
         "ScripTree",
         "Works with any executable. No source modification. GUI defined "
         "externally in a portable JSON file."],
    ]
    add_styled_table(doc, cat_headers, cat_rows, col_widths=[4, 4, 9])

    # ================================================================
    # 4. FEATURE COMPARISON MATRIX
    # ================================================================
    doc.add_heading("4. Feature Comparison Matrix", level=1)
    doc.add_paragraph(
        "The following matrix compares capabilities across all surveyed tools. "
        "A check mark (\u2713) indicates full support; a tilde (~) indicates "
        "partial or limited support; a dash (\u2014) indicates no support."
    )

    feat_headers = [
        "Feature", "ScripTree", "Gooey", "Fire", "Clicky",
        "Sunbeam", "Gradio",
    ]
    feat_rows = [
        ["Works with any executable",
         "\u2713", "\u2014", "\u2014", "\u2014", "~", "\u2014"],
        ["No source code changes needed",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Auto-parse --help output",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Blank-canvas manual editor",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Native OS file/folder dialogs",
         "\u2713", "\u2713", "\u2014", "\u2014", "\u2014", "~"],
        ["Portable definition files",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Tool tree / launcher",
         "\u2713", "\u2014", "\u2014", "\u2014", "~", "\u2014"],
        ["Named configurations / presets",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Sections / tabs for complex tools",
         "\u2713", "~", "\u2014", "\u2014", "\u2014", "~"],
        ["Real-time output streaming",
         "\u2713", "\u2713", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Environment & PATH overrides",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["LLM-ready documentation",
         "\u2713", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014"],
        ["Cross-language support",
         "\u2713", "\u2014", "\u2014", "\u2014", "~", "\u2014"],
        ["Works offline (no browser)",
         "\u2713", "\u2713", "\u2713", "\u2713", "\u2713", "\u2014"],
        ["Python-only requirement",
         "\u2014", "\u2713", "\u2713", "\u2713", "\u2014", "\u2713"],
        ["Web-based UI",
         "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "\u2713"],
    ]
    add_styled_table(doc, feat_headers, feat_rows,
                     col_widths=[5, 2, 1.6, 1.3, 1.5, 1.8, 1.6])

    # ================================================================
    # 5. DETAILED COMPETITOR PROFILES
    # ================================================================
    doc.add_heading("5. Detailed Competitor Profiles", level=1)

    # --- Gooey ---
    doc.add_heading("5.1 Gooey", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Website: ")
    run.bold = True
    p.add_run("github.com/chriskiehl/Gooey")

    doc.add_paragraph(
        "Gooey is the closest competitor in concept. It converts Python "
        "argparse-based CLI programs into GUI applications by adding a "
        "@Gooey decorator to the script's main function. At runtime, it "
        "introspects the ArgumentParser object and generates a wxPython "
        "form with appropriate widgets."
    )

    doc.add_heading("Strengths", level=3)
    for s in [
        "Mature project (10+ years, 20k+ GitHub stars).",
        "Excellent argparse integration \u2014 if your tool uses argparse, "
        "the GUI is nearly automatic.",
        "Supports file/folder choosers, date pickers, progress bars.",
        "Layout customisation via GooeyParser subgroups.",
        "Output panel with real-time streaming.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Limitations vs. ScripTree", level=3)
    for s in [
        "Python-only: cannot wrap C#, C++, Rust, Go, or any compiled "
        "executable.",
        "Requires source modification: the @Gooey decorator and import "
        "must be added to the script. You need access to the source code.",
        "No external definition file: the GUI is baked into the script, "
        "not portable or shareable separately.",
        "No tool launcher/tree: each Gooey script is its own standalone "
        "application.",
        "No blank-canvas workflow: if the tool doesn't use argparse, "
        "Gooey has no path forward.",
        "wxPython dependency can be difficult to install on some systems.",
        "No configurations/presets system.",
        "Maintenance has slowed \u2014 last significant release was 2020.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # --- Python Fire ---
    doc.add_heading("5.2 Python Fire (Google)", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Website: ")
    run.bold = True
    p.add_run("github.com/google/python-fire")

    doc.add_paragraph(
        "Python Fire automatically generates a CLI from any Python object "
        "\u2014 functions, classes, modules, or even dictionaries. It "
        "introspects type hints and docstrings to build help text and "
        "argument parsing."
    )

    doc.add_heading("Strengths", level=3)
    for s in [
        "Zero-effort CLI generation from existing Python code.",
        "Excellent for developer tooling and REPL exploration.",
        "Google-backed, well-maintained.",
        "Supports nested commands from class hierarchies.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Limitations vs. ScripTree", level=3)
    for s in [
        "Generates a CLI, not a GUI. There is no visual interface at all.",
        "Python-only: cannot work with non-Python executables.",
        "Requires importing Fire and calling fire.Fire() in the source.",
        "No file pickers, no dropdowns, no form layout.",
        "Different problem space entirely \u2014 Fire makes CLIs easier to "
        "write, not easier to use.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # --- Clicky ---
    doc.add_heading("5.3 Clicky", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Website: ")
    run.bold = True
    p.add_run("github.com/pyrustic/clicky")

    doc.add_paragraph(
        "Clicky is a newer Python library that generates a CLI and "
        "optional TUI from decorated Python functions. It focuses on "
        "rapid prototyping and uses type annotations to infer argument "
        "types."
    )

    doc.add_heading("Strengths", level=3)
    for s in [
        "Clean, modern API with decorator syntax.",
        "Automatic type coercion from annotations.",
        "Integrated help generation.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Limitations vs. ScripTree", level=3)
    for s in [
        "Python-only, requires source modification.",
        "Generates terminal UI, not a native GUI with OS dialogs.",
        "No file pickers or native widgets.",
        "Very early stage \u2014 small community, limited documentation.",
        "No external tool definitions or launcher system.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # --- Sunbeam ---
    doc.add_heading("5.4 Sunbeam", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Website: ")
    run.bold = True
    p.add_run("github.com/pomdtr/sunbeam")

    doc.add_paragraph(
        "Sunbeam is a command launcher inspired by Raycast and "
        "fzf. Extensions are scripts (any language) that read JSON on "
        "stdin and write JSON describing UI pages to stdout. The renderer "
        "is a terminal-based fuzzy-finder interface."
    )

    doc.add_heading("Strengths", level=3)
    for s in [
        "Language-agnostic extensions \u2014 any script that speaks JSON.",
        "Fast, keyboard-driven launcher UI.",
        "Composable: extensions can chain into each other.",
        "Growing ecosystem of community extensions.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Limitations vs. ScripTree", level=3)
    for s in [
        "Every tool must be written as a Sunbeam extension \u2014 existing "
        "executables don't work out of the box.",
        "Terminal-only UI: no native file dialogs, no rich form widgets.",
        "Designed as a launcher, not a form builder \u2014 poor fit for "
        "tools with many parameters.",
        "No parameter type system (no integer validation, no enums, "
        "no multiselect).",
        "No saved configurations or presets.",
        "Primarily macOS/Linux \u2014 Windows support is secondary.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # --- Gradio ---
    doc.add_heading("5.5 Gradio", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Website: ")
    run.bold = True
    p.add_run("gradio.app")

    doc.add_paragraph(
        "Gradio creates web-based UIs for Python functions, primarily "
        "targeting machine-learning demos. It supports a wide range of "
        "input/output components and can be shared via public URLs."
    )

    doc.add_heading("Strengths", level=3)
    for s in [
        "Rich widget library (sliders, image inputs, audio, code blocks).",
        "One-click public sharing via Hugging Face Spaces.",
        "Excellent for ML model demos and data science workflows.",
        "Active development with large community.",
        "Built-in API generation alongside the UI.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Limitations vs. ScripTree", level=3)
    for s in [
        "Python-only: requires wrapping functionality in Python functions.",
        "Browser-based: requires a running web server, not a native app.",
        "Cannot wrap arbitrary executables \u2014 only Python callables.",
        "No portable definition files; the UI is defined in Python code.",
        "Overkill for simple CLI tools; designed for ML/data workflows.",
        "No offline-first desktop experience.",
        "No tool tree or launcher concept.",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # --- Adjacent tools ---
    doc.add_heading("5.6 Adjacent tools (not direct competitors)", level=2)

    doc.add_paragraph(
        "Several other tools occupy nearby spaces but solve fundamentally "
        "different problems:"
    )

    adj_headers = ["Tool", "What it does", "Why it's not a competitor"]
    adj_rows = [
        ["Streamlit",
         "Web framework for data apps in Python.",
         "Requires writing a Python app from scratch. Not a wrapper "
         "for existing tools."],
        ["Textual (Textualize)",
         "Python framework for rich terminal UIs.",
         "A TUI toolkit, not a GUI generator. Developer builds the "
         "interface in code."],
        ["Charm / Bubbletea",
         "Go framework for terminal UIs.",
         "Same category as Textual. Manual TUI development, not "
         "auto-generation."],
        ["Wails / Tauri",
         "Desktop app frameworks (Go/Rust + web frontend).",
         "General-purpose app frameworks. Building a CLI wrapper would "
         "require writing the entire UI."],
        ["Zenity / YAD",
         "Simple GTK dialog boxes from shell scripts.",
         "One-shot dialogs only (file picker, message box). No form "
         "builder, no parameter system, no tool definitions."],
    ]
    add_styled_table(doc, adj_headers, adj_rows, col_widths=[3, 5.5, 8.5])

    # ================================================================
    # 6. SCRIPTREE'S UNIQUE ADVANTAGES
    # ================================================================
    doc.add_heading("6. ScripTree\u2019s Unique Advantages", level=1)
    doc.add_paragraph(
        "No other tool in the survey matches ScripTree on all of the "
        "following dimensions simultaneously:"
    )

    doc.add_heading("6.1 Language-agnostic, source-agnostic", level=2)
    doc.add_paragraph(
        "ScripTree wraps any executable: Python scripts, compiled C# "
        "applications, Rust CLIs, legacy batch files, system utilities "
        "like robocopy or ffmpeg. The user never needs access to or "
        "knowledge of the tool's source code. This is the single biggest "
        "differentiator \u2014 every other tool requires either Python "
        "source modification or writing tool-specific adapters."
    )

    doc.add_heading("6.2 Dual ingest: auto-parse + blank canvas", level=2)
    doc.add_paragraph(
        "ScripTree offers two equally supported paths to creating a tool "
        "definition. For tools with standard --help output, the parser "
        "detects argument patterns and assigns appropriate widgets "
        "automatically. For tools with no help text (common with custom "
        "executables), the blank-canvas editor lets users define "
        "parameters manually in under 90 seconds. No competitor offers "
        "both paths."
    )

    doc.add_heading("6.3 Portable JSON definitions", level=2)
    doc.add_paragraph(
        "A .scriptree file is a self-contained JSON document that can be "
        "version-controlled, shared via email, bundled with the tool it "
        "wraps, or generated programmatically. The GUI definition lives "
        "outside the tool, not inside it. This enables workflows "
        "impossible with decorator-based tools: a sysadmin creates "
        ".scriptree files for system utilities and distributes them to the "
        "team without touching the utilities themselves."
    )

    doc.add_heading("6.4 Tool tree launcher", level=2)
    doc.add_paragraph(
        "The .scriptreetree format groups multiple tools into a navigable "
        "hierarchy \u2014 an entire team's toolkit in one window. This "
        "transforms ScripTree from a single-tool wrapper into a daily-"
        "driver application. No competitor offers an equivalent."
    )

    doc.add_heading("6.5 Configurations and environment management", level=2)
    doc.add_paragraph(
        "Each tool can have multiple named configurations (parameter "
        "presets), each with its own environment variable overrides and "
        "PATH modifications. This handles real-world scenarios like "
        "\"development vs. production\" or \"local vs. remote\" without "
        "duplicating tool definitions."
    )

    doc.add_heading("6.6 Sections and tabs for complex tools", level=2)
    doc.add_paragraph(
        "Tools with many parameters (robocopy has 60+) can organise them "
        "into collapsible sections or tabbed pages. The user sees a clean, "
        "categorised form instead of an overwhelming flat list. Gooey "
        "offers basic grouping; no other competitor matches this level of "
        "layout control."
    )

    # ================================================================
    # 7. LLM-POWERED WORKFLOW
    # ================================================================
    doc.add_heading("7. LLM-Powered Workflow", level=1)
    doc.add_paragraph(
        "ScripTree includes a dedicated documentation pack designed "
        "specifically for AI assistants. This is a unique capability that "
        "no competitor offers."
    )

    doc.add_heading("7.1 The LLM help file suite", level=2)
    doc.add_paragraph(
        "Located in the help/LLM/ directory, these files provide "
        "everything an AI agent needs to generate valid .scriptree files:"
    )

    llm_headers = ["File", "Purpose"]
    llm_rows = [
        ["scriptree_format.md",
         "Complete JSON schema for .scriptree files with every field, "
         "type, and loader invariant documented."],
        ["scriptreetree_format.md",
         "Tree launcher format and path resolution rules."],
        ["argument_template.md",
         "The substitution grammar for argv assembly, including "
         "conditional flags, groups, and edge cases."],
        ["param_types_widgets.md",
         "Type-to-widget compatibility matrix, default values per type, "
         "and coercion rules."],
        ["configurations_sidecar.md",
         "Sidecar JSON format for named configurations with env/PATH "
         "overrides."],
        ["architecture.md",
         "Package layout and the core vs. UI separation."],
        ["parsers/python_scripts.md",
         "Rules for generating Python CLIs whose --help imports cleanly."],
        ["parsers/windows_exe.md",
         "Rules for Windows executables and their flag conventions."],
        ["parsers/gnu_tools.md",
         "Rules for GNU/Linux tools and their --help patterns."],
    ]
    add_styled_table(doc, llm_headers, llm_rows, col_widths=[5, 12])

    doc.add_heading("7.2 How LLM generation works", level=2)

    steps = [
        ("User describes the tool", "\"I have a C# exe that takes a "
         "source folder, output folder, and a --format flag with options "
         "json/xml/csv.\""),
        ("AI reads the LLM docs", "The agent reads scriptree_format.md "
         "and param_types_widgets.md to understand the schema."),
        ("AI generates the .scriptree file", "Produces valid JSON with "
         "correct types, widgets, argument template, and section layout."),
        ("User opens in ScripTree", "The generated file loads directly "
         "\u2014 no manual cleanup needed."),
    ]
    for i, (step, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Step {i} \u2014 {step}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph("")
    doc.add_paragraph(
        "This workflow means users don't even need to open the ScripTree "
        "editor for straightforward tools. They describe what they want "
        "in natural language, and the AI produces a ready-to-use "
        "definition file. This dramatically lowers the barrier to entry "
        "and is especially powerful for teams that use AI coding "
        "assistants like Claude Code, GitHub Copilot, or ChatGPT."
    )

    # ================================================================
    # 8. MARKET POSITIONING
    # ================================================================
    doc.add_heading("8. Market Positioning & Target Audiences", level=1)

    doc.add_heading("8.1 Primary audiences", level=2)

    audiences = [
        ("Sysadmins & DevOps engineers",
         "Run the same robocopy, rsync, ffmpeg, docker commands daily "
         "with varying arguments. ScripTree turns tribal knowledge into "
         "clickable forms that can be handed to junior staff. The "
         "configurations feature maps perfectly to dev/staging/prod "
         "environments."),
        ("Technical non-developers",
         "CAD engineers, data analysts, lab technicians, and video "
         "editors who have powerful CLI tools installed but avoid them "
         "because the learning curve feels steep. A file picker and a "
         "few dropdowns removes the entire barrier."),
        ("Team leads who need to delegate",
         "\"Here, run this tool with these settings\" becomes a "
         ".scriptree file instead of a wiki page with screenshots that's "
         "always out of date. The tree launcher makes it a one-stop shop."),
        ("Script authors & internal tool builders",
         "Developers who write Python/PowerShell utilities for their "
         "team. Today they either build a basic Tkinter GUI (hours of "
         "work) or tell people to use the command line (hours of "
         "support). ScripTree is the missing middle: describe the params, "
         "get a real GUI, done."),
        ("AI-assisted development teams",
         "Teams using AI coding assistants can have the AI generate "
         ".scriptree files as part of their tool-building workflow. The "
         "LLM documentation pack makes this reliable and repeatable."),
    ]
    for title_text, desc in audiences:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("8.2 Positioning statement", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    run = p.add_run(
        "\"ScripTree turns any command-line tool into a native Windows "
        "application \u2014 no coding required, no source access needed. "
        "Point it at an executable, define the parameters, and get a "
        "professional form with file pickers, dropdowns, and saved "
        "presets. Build once, share as a portable JSON file, organise "
        "your tools into a launcher, and let AI generate new tool "
        "definitions for you.\""
    )
    run.italic = True

    # ================================================================
    # 9. MASS APPEAL ASSESSMENT
    # ================================================================
    doc.add_heading("9. Mass Appeal Assessment", level=1)

    doc.add_heading("9.1 Growth model", level=2)
    doc.add_paragraph(
        "ScripTree's adoption pattern would mirror tools like WinSCP, "
        "Notepad++, and 7-Zip: unglamorous, genuinely useful, and "
        "impossible to stop once one person introduces it to a team. "
        "The predicted growth model is workplace-viral rather than "
        "social-media-viral."
    )

    doc.add_heading("9.2 Viral mechanics", level=2)
    viral = [
        ("One-to-many sharing", "A single power user creates .scriptree "
         "files and shares them with the team. Recipients need only "
         "ScripTree installed \u2014 they never touch the editor."),
        ("Tree launcher as daily driver", "Once a user has a curated tree "
         "of their 15 most-used tools, ScripTree becomes indispensable. "
         "The switching cost is high because the value is cumulative."),
        ("AI amplification", "AI assistants can generate .scriptree files "
         "from --help output or natural-language descriptions, making "
         "tool creation nearly effortless. This lowers the barrier to the "
         "\"create\" side, not just the \"use\" side."),
    ]
    for title_text, desc in viral:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("9.3 Adoption barriers", level=2)
    barriers = [
        ("Discovery problem", "Most people don't search for \"GUI wrapper "
         "for CLI tools\" because they don't know the category exists. "
         "Marketing needs to target specific pain points: \"tired of "
         "remembering ffmpeg flags?\" or \"make your Python scripts "
         "team-friendly.\""),
        ("Power-user paradox", "The people who most need this (CLI-averse "
         "users) don't know the tools exist. The people who know the "
         "tools exist (power users) often think they don't need a GUI. "
         "The sweet spot is power users building forms for others."),
        ("Platform scope", "Windows-first is smart for the initial "
         "audience (sysadmins, corporate, CAD), but broad adoption "
         "requires cross-platform. The Qt/PySide6 foundation makes this "
         "achievable with the existing core/ui separation."),
    ]
    for title_text, desc in barriers:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(title_text + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("9.4 Verdict", level=2)
    doc.add_paragraph(
        "ScripTree has strong potential for organic, workplace-driven "
        "adoption. It fills a genuine gap that no existing tool addresses: "
        "wrapping arbitrary executables in a native GUI without touching "
        "source code. The .scriptree file format is its secret weapon \u2014 "
        "portable, version-controllable, AI-generatable, and shareable. "
        "The tree launcher is the retention hook that turns a utility "
        "into a daily driver."
    )
    doc.add_paragraph(
        "The most likely growth trajectory is not \"go viral on Hacker "
        "News\" but rather \"quietly becomes indispensable in thousands of "
        "workplaces\" \u2014 which is, historically, the more durable and "
        "valuable kind of adoption."
    )

    # ================================================================
    # 10. CONCLUSION
    # ================================================================
    doc.add_heading("10. Conclusion", level=1)

    doc.add_paragraph(
        "ScripTree occupies a unique position in the CLI-to-GUI space. "
        "It is the only tool surveyed that:"
    )
    for item in [
        "Works with any executable regardless of language or framework.",
        "Requires no source code modification.",
        "Auto-parses --help output and supports blank-canvas manual "
        "definition equally.",
        "Saves GUI definitions as portable, shareable JSON files.",
        "Groups tools into navigable tree launchers.",
        "Supports named configurations with environment overrides.",
        "Provides LLM-ready documentation for AI-generated definitions.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("")
    doc.add_paragraph(
        "The competitive landscape validates that demand exists (Gooey's "
        "20,000+ GitHub stars prove people want CLI-to-GUI tooling) but "
        "that current solutions are limited to Python-only, source-"
        "modification-required approaches. ScripTree removes both "
        "constraints while adding capabilities (tree launcher, "
        "configurations, LLM generation) that no competitor matches."
    )

    return doc


if __name__ == "__main__":
    out_dir = Path(r"C:\Users\Ken\OneDrive\Kens_Projects\Claude\Software\ScripTree\docs")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "ScripTree_Competitive_Analysis.docx"

    doc = build_document()
    doc.save(str(out_path))
    print(f"Saved: {out_path}")
